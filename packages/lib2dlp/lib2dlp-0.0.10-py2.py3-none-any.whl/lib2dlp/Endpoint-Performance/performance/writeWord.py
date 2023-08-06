#!/usr/bin/python
# -*- coding: UTF-8 -*-

import docx
import time
import random
import shutil
import string
import difflib
from faker import Factory
from faker import Faker
from docx import Document
from docx.shared import Inches
import os
import sys
reload(sys)
sys.setdefaultencoding('utf-8')

def genStr(num):
   H = '0123456789ABCDEFGHIJKLMNOPQISTUVWXYZabcdefghijklmnopqistuvwxyz'
   salt = ''
   for i in range(num):
      salt += random.choice(H)
   return salt

def timeget():
    timeStamp = time.time()
    timeArray = time.localtime(timeStamp)
    otherStyleTime = time.strftime("%Y-%m-%d %H:%M:%S", timeArray)
    return otherStyleTime

def string_similar(a, b):
   return difflib.SequenceMatcher(None, a, b).quick_ratio()

def findKeyWord(full_path):
   # 首先将doc转换成docx
   word = wc.Dispatch("Word.Application")

   # 找到word路径 + 文件名 ，即可打开文件
   #full_path = 'D:\\performance2\\data\\test.docx'
   doc1 = word.Documents.Open(full_path)

   # 读取word内容
   #　这里是以段落为单位的，下面用一个for 遍历所有段落

   doc = docx.Document(full_path)
   #parag_num = 0
   for para in doc.paragraphs :
      #print(para.text)
      similar = string_similar(para.text,u'拍照点阵水印')
      #print similar
   doc1.Close()
   if similar == 1:
      print "文件打开成功"
      return 1
   else:
      print "文件打开失败"
      return 0

def CreateWord(s):
   faker = Faker("zh_CN")
   i = 0
   z = 0
   document = Document()

   while i < s:
      document.add_heading(u'文档标题', 0)
      p = document.add_paragraph(u'这是一个自然段 ')
      p.add_run('bold').bold = True
      p.add_run(u' 还有 ')
      p.add_run('italic.').italic = True

      document.add_heading(u'1级别标题', level=1)
      document.add_paragraph(u'引用', style='IntenseQuote')

      document.add_paragraph(u'符号列表', style='ListBullet')
      document.add_paragraph(u'数字列表t', style='ListNumber')
      document.add_paragraph(u'我的微信:')
      document.add_picture('data\\tupian.jpg', width=Inches(3.25))

      table = document.add_table(rows=3, cols=3)
      hdr_cells = table.rows[0].cells
      hdr_cells[0].text = u'第一列'
      hdr_cells[1].text = u'第二列'
      hdr_cells[2].text = u'第三列'

      hdr_cells = table.rows[1].cells
      hdr_cells[0].text = '1'
      hdr_cells[1].text = '21'
      hdr_cells[2].text = 'qwertyuiop'

      hdr_cells = table.rows[2].cells
      hdr_cells[0].text = '2'
      hdr_cells[1].text = '43'
      hdr_cells[2].text = 'asdfghjkl'
      i = i + 1

   document.add_heading(u'我是关键字', 0)
   while z < 10:
      #salt = ''.join(random.sample(string.ascii_letters + string.digits, 16))
      document.add_paragraph(u'参考资料' + genStr(32))
      z = z + 1

   x = 1
   docText = faker.text(max_nb_chars=500*2048)
   while x < 11:
      document.add_paragraph(docText)
      x = x + 1

   document.add_page_break()

   document.save("data\\write\\dlp_%s.docx" % genStr(8))

   return True

if __name__ == "__main__":
   while 1:
      if os.path.exists("data\\write"):
         shutil.rmtree("data\\write",ignore_errors=True)
         os.mkdir("data\\write")
         print u"——>\t先删除后创建"
      else:
         os.mkdir("data\\write")
         print u"——>\t直接创建"
      s = random.randint(1,10) #如果是10约10M
      CreateWord(s)
      print u"%sM——>\t%s" % (s,timeget())
      time.sleep(3)