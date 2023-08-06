#!/usr/bin/python
# -*- coding: UTF-8 -*-
#载入必要的模块

from datetime import datetime
from faker import Factory
from faker import Faker
from docx import Document
from docx.shared import Inches
import time
import random
import sys
import os
import shutil
import ConfigParser
import string


reload(sys)
sys.setdefaultencoding('utf8')

class office:

    def hello(self):
        print('hello')

    def timeget(self):
        timeStamp = time.time()
        timeArray = time.localtime(timeStamp)
        otherStyleTime = time.strftime("%Y%m%d%H%M%S", timeArray) 
        return otherStyleTime

    def mkdir(self,path):
        path=path.strip()
        path=path.rstrip("\\")
        isExists=os.path.exists(path)
        if not isExists:
            os.makedirs(path.decode('utf-8'))
            print path + u' 创建成功'
            return path
        else:
            print path + u' 目录已存在'
            return path

    def CreateWordNum(self,num):
        path = "D:\\performance2\\Result\\word"
        o.mkdir(path)
        i = 0
        z = 0
        document = Document()
        while i < 100:
            document.add_heading(u'文档标题', 0)
            p = document.add_paragraph(u'这是一个自然段 ')
            p.add_run('bold').bold = True
            p.add_run(u' 还有 ')
            p.add_run('italic.').italic = True

            document.add_heading(u'1级别标题', level=1)
            document.add_paragraph(u'引用', style='IntenseQuote')

            document.add_paragraph(u'符号列表', style='ListBullet')
            document.add_paragraph(u'数字列表t', style='ListNumber')
            document.add_paragraph(u'我的微信:')
            document.add_picture('data\\tupian.jpg', width=Inches(3.25))

            table = document.add_table(rows=3, cols=3)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = u'第一列'
            hdr_cells[1].text = u'第二列'
            hdr_cells[2].text = u'第三列'

            hdr_cells = table.rows[1].cells
            hdr_cells[0].text = '1'
            hdr_cells[1].text = '21'
            hdr_cells[2].text = 'qwertyuiop'

            hdr_cells = table.rows[2].cells
            hdr_cells[0].text = '2'
            hdr_cells[1].text = '43'
            hdr_cells[2].text = 'asdfghjkl'
            i = i + 1

        document.add_heading(u'我是关键字', 0)
        s = random.randint(0,50)#之前是50
        while z < s:
            faker = Faker("zh_CN")
            text = faker.text(max_nb_chars=20480)#之前是20480
            salt = ''.join(random.sample(string.ascii_letters + string.digits, 16))
            document.add_paragraph(u'参考资料' + text + salt)
            z = z + 1

        document.add_page_break()

        document.save("D:\\performance2\\Result\\word\\dlp%s.docx" % num)

        return True
    def CreateWordE(self,n):
        m = 0
        o = office()
        while m < n:
            o.CreateWordNum(m)
            m = m + 1
        return True
    def CreateWordE1(self,n):
        m = 0
        o = office()
        while m < n:
            o.CreateWordNum(m)
            m = m + 1
        return True

    
    def get_desktop(self):
        key = _winreg.OpenKey(_winreg.HKEY_CURRENT_USER,r'Software\Microsoft\Windows\CurrentVersion\Explorer\Shell Folders')
        return _winreg.QueryValueEx(key, "Desktop")[0]



    def logtodis(self,text):
        logpath = 'D:\\performance2\\log\\rundiscovery.txt'
        f = open(logpath,'a+')
        f.write(text)
        f.write('---rundiscovery:\t%s\r\n' %datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        f.close()

    def rundiscovery(self):
        o = office()
        path = "D:\\performance2\\Result\\word"
        isExists = os.path.exists(path)
        if not isExists:
            o.mkdir(path)
        else:
            shutil.rmtree(path)
            time.sleep(3)
            o.mkdir(path)
        
        print "START------>" + datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        o.CreateWordE1(9)#之前是9
        time.sleep(5)
        o.logtodis("WORD")
        
        print "STOP------>" + datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    def test(self):
        i = 0
        path = os.getcwd()+ '\\'
        time.sleep(3)
        starttime = time.time()
        for y in range(18):
            try:
                o = office()
                o.rundiscovery()
                time.sleep(3)
                f = o.timeget()
                print 'rundiscovery' + '  '+ str(i)+ '  ' + f

                i = i + 1
            except Exception,ex:
                print ex
            if y == 17:
                break
            time.sleep(3600)
        print 'over'
        win32api.keybd_event(win32con.VK_LWIN,0,0,0)
        win32api.keybd_event(68,0,0,0)
        win32api.keybd_event(68,0,win32con.KEYEVENTF_KEYUP,0)
        win32api.keybd_event(win32con.VK_LWIN,0,win32con.KEYEVENTF_KEYUP,0)          
        subprocess.Popen('taskmgr')
        time.sleep(4)
        im = ImageGrab.grab()
        im.save('c:\\12.png')      
        time.sleep(120)
        try:
            os.system('taskkill /IM 360tray.exe /F')
            os.system('taskkill /IM 360EntClient.exe /F')
            os.system('taskkill /IM DLPMain.exe /F')
            os.system('taskkill /IM DlpCenter.exe /F')
            os.system('taskkill /IM dlpmain32.exe /F')
            os.system('taskkill /IM DriverControler.exe /F')
            os.system('taskkill /IM DriverControler.exe /F')
        except Exception,ex:
            print ex
    
#end--------------------------------------------------------

if __name__ == "__main__":
    o = office()
    #o.test()
    o.rundiscovery()