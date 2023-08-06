#!/usr/bin/python
# -*- coding: UTF-8 -*-

from win32com import client as wc
from faker import Factory
from faker import Faker
from collections import Counter
from docx import Document
from docx.shared import Inches
from PIL import Image
from PIL import ImageGrab
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.image import MIMEImage
from email.header import Header
from email.utils import parseaddr, formataddr
import docx
import matplotlib.pyplot as plt
import os
import difflib
import random
import string
import win32gui
import win32ui
import win32api
import win32con
import time
import smtplib
import email
import ctypes
import sys
reload(sys)
sys.setdefaultencoding("utf-8")


def string_similar(a, b):
   return difflib.SequenceMatcher(None, a, b).quick_ratio()

def CreateWord(s):
   faker = Faker("zh_CN")
   i = 0
   z = 0
   document = Document()

   while i < s:
      document.add_heading(u'文档标题', 0)
      p = document.add_paragraph(u'这是一个自然段 ')
      p.add_run('bold').bold = True
      p.add_run(u' 还有 ')
      p.add_run('italic.').italic = True

      document.add_heading(u'1级别标题', level=1)
      document.add_paragraph(u'引用', style='IntenseQuote')

      document.add_paragraph(u'符号列表', style='ListBullet')
      document.add_paragraph(u'数字列表t', style='ListNumber')
      document.add_paragraph(u'我的微信:')
      document.add_picture('data\\tupian.jpg', width=Inches(3.25))

      table = document.add_table(rows=3, cols=3)
      hdr_cells = table.rows[0].cells
      hdr_cells[0].text = u'第一列'
      hdr_cells[1].text = u'第二列'
      hdr_cells[2].text = u'第三列'

      hdr_cells = table.rows[1].cells
      hdr_cells[0].text = '1'
      hdr_cells[1].text = '21'
      hdr_cells[2].text = 'qwertyuiop'

      hdr_cells = table.rows[2].cells
      hdr_cells[0].text = '2'
      hdr_cells[1].text = '43'
      hdr_cells[2].text = 'asdfghjkl'
      i = i + 1

   document.add_heading(u'我是关键字', 0)
   while z < 10:
      salt = ''.join(random.sample(string.ascii_letters + string.digits, 16))
      document.add_paragraph(u'王宝强' + salt)
      z = z + 1

   x = 1
   docText = faker.text(max_nb_chars=500*2048)
   while x < 11:
      document.add_paragraph(docText)
      x = x + 1

   document.add_page_break()

   document.save("Result\\dlp.docx")

   return True

def window_capture():
   hwnd = 0
   dpath = 'Result\\'
   hwndDC = win32gui.GetWindowDC(hwnd)
   mfcDC=win32ui.CreateDCFromHandle(hwndDC)
   saveDC=mfcDC.CreateCompatibleDC()
   saveBitMap = win32ui.CreateBitmap()
   MoniterDev=win32api.EnumDisplayMonitors(None,None)
   w = MoniterDev[0][2][2]
   h = MoniterDev[0][2][3]
   #print w,h　　　＃图片大小
   saveBitMap.CreateCompatibleBitmap(mfcDC, w, h)
   saveDC.SelectObject(saveBitMap)
   saveDC.BitBlt((0,0),(w, h) , mfcDC, (0,0), win32con.SRCCOPY)
   cc=time.gmtime()
   bmpname=str(cc[0])+str(cc[1])+str(cc[2])+str(cc[3]+8)+str(cc[4])+str(cc[5])+'_WindowsCapture'+'.bmp'
   saveBitMap.SaveBitmapFile(saveDC, bmpname)
   Image.open(bmpname).save(bmpname[:-4]+".jpg")
   os.remove(bmpname)
   jpgname=bmpname[:-4]+'.jpg'
   djpgname=dpath+jpgname  
   copy_command = "move %s %s" % (jpgname, djpgname)
   os.popen(copy_command)
   return True

def addAttch(path):
   faker = Faker("zh_CN")
   SUBJECT = '%s发送的邮件' % faker.name()
   FROM = 'datasec360@126.com'
   PASSWORD = '1234qwer'
   To = 'datasec360@163.com'
   msg = MIMEMultipart('related')

   # 创建一个用于发送文本的MIMEText对象
   content = faker.text(max_nb_chars=2048)
   msg_text = MIMEText(content,_subtype='plain',_charset='utf-8')
   msg.attach(msg_text)

   os.chdir(path)
   dir = os.getcwd()

   # 将xls作为附件添加到邮件中
   # 创建MIMEText对象，保存txt文件
   for fn in os.listdir(dir):
      attach = MIMEText(open(fn,'rb').read())
      # 指定当前文件格式类型
      attach['Content-type'] = 'application/octet-stream'
      # 配置附件显示的文件名称,当点击下载附件时，默认使用的保存文件的名称
      # gb18030 qq邮箱中使用的是gb18030编码，防止出现中文乱码
      #attach['Content-Disposition'] = 'attachment;filename="我是tupian附件.txt"'.decode('utf-8').encode('gb18030')
      attach['Content-Disposition'] = 'attachment;filename=' + fn
      # 把附件添加到msg中
      msg.attach(attach)
   # 设置必要请求头信息
   msg['From'] = FROM
   msg['To'] = To
   msg['Subject'] = SUBJECT

   return msg

def sendMail(msg):
   HOST = 'smtp.126.com'
   FROM = 'datasec360@126.com'
   PASSWORD = '1234qwer'
   To = 'datasec360@163.com'
   ret = True
   try:
      # 发送邮件
      server = smtplib.SMTP() # SMTP协议默认端口是25
      server.connect(HOST, 25)
      #server.ehlo() 
      #server.starttls()  #开启加密传输
      #server.set_debuglevel(0)#打印出和SMTP服务器交互的所有信息，不想打印就设置为0
      server.login(FROM,PASSWORD)#登录服务器
      server.sendmail(FROM,To,msg.as_string())
      server.quit()
      print "SMTP OK" 
      return ret
   except Exception,e:
      print e
      print "SMTP False" 
      return False
   return

def run():
   Winmm = ctypes.WinDLL('Winmm.dll')
   # 首先将doc转换成docx
   word = wc.Dispatch("Word.Application")
   b_t = Winmm.timeGetTime()

    
   # 找到word路径 + 文件名 ，即可打开文件
   full_path = 'D:\\performance2\\data\\test.docx'
   doc1 = word.Documents.Open(full_path)

   # 读取word内容
   #　这里是以段落为单位的，下面用一个for 遍历所有段落

   doc = docx.Document("data\\test.docx")
   #parag_num = 0
   for para in doc.paragraphs :
      #print(para.text)
      similar = string_similar(para.text,u'拍照点阵水印')
      #print similar
   ret = Winmm.timeGetTime() - b_t
   print(ret)
   doc1.Close()
   if similar == 1:
      print "文件打开成功"
      return 1
   else:
      print "文件打开失败"
      return 0

if __name__ == "__main__":
   #s = 10
   #path = "Result"
   run()
   #CreateWord(s)

   #window_capture()
   
   #msg = addAttch(path)
   #sendMail(msg)