#!/usr/bin/python
# -*- coding: UTF-8 -*-

from faker import Factory
from faker import Faker
from docx import Document
from docx.shared import Inches
import os.path
import random
import string
import time

keyword1 = ["十六开本","大众市场纸皮书版式","启蒙读物","文摘通报","学位论文","组稿编辑","胶粘装订联动线","青少年读物","文稿代理人","出版日期和书价"]
keyword2 = ["会议建议","会议指出","持续进行","会议充分肯定了","正式启动","会议安排","会议倡议","会议要求贯彻落实","各部门依照","整个会议共持续"]
keyword3 = ["政治主张","自治请愿","支持新疆独立","恐怖组织","勾结外国","打倒资本主义","基督教民主主义","权威主义","公共言论的抑制","言论自由受到"]
lang_list = ["ar_AA", "ar_EG", "ar_JO", "ar_PS", "ar_SA", "bg_BG", "bs_BA", "cs_CZ", "de", "de_AT", "de_CH", "de_DE", "dk_DK", "el_CY", "el_GR", "en", "en_AU", "en_CA", "en_GB", "en_IE", "en_NZ", "en_TH", "en_US", "es", "es_ES", "es_MX", "et_EE", "fa_IR", "fi_FI", "fr_CH", "fr_FR", "he_IL", "hi_IN", "hr_HR", "hu_HU", "id_ID", "it_IT", "ja_JP", "ka_GE", "ko_KR", "la", "lb_LU", "lt_LT", "lv_LV", "mt_MT", "ne_NP", "nl_BE", "nl_NL", "no_NO", "pl_PL", "pt_BR", "pt_PT", "ro_RO", "ru_RU", "sk_SK", "sl_SI", "sv_SE", "th_TH", "tr_TR", "tw_GH", "uk_UA", "zh_CN", "zh_TW"]


def makedir(path):
   path=path.strip()
   path=path.rstrip("\\")
   isExists=os.path.exists(path)
   if not isExists:
      os.makedirs(path.decode('utf-8'))
      #print path + u' 创建成功'
      return path
   else:
      #print path + u' 目录已存在'
      return path

def matchlanguage(lang):
   if lang == "zh_CN":
      return u"简体中文-中国"
   elif lang == "zh_TW":
      return u"繁体中文-中国台湾"
   elif lang == "en_US":
      return u"英文-美国"
   elif lang == "en_GB":
      return u"英文-英国"
   elif lang == "de_DE":
      return u"德文-德国"
   elif lang == "ja_JP":
      return u"日文-日本"
   elif lang == "ko_KR":
      return u"朝鲜文-韩国"
   elif lang == "fr_FR":
      return u"法文-法国"
   elif lang == "ar_AA":
      return u"阿拉伯语"
   elif lang == "ar_EG":
      return u"阿拉伯语"
   elif lang == "ar_PS":
      return u"阿拉伯语-巴勒斯坦"
   elif lang == "ar_SA":
      return u"阿拉伯语-沙特阿拉伯"
   elif lang == "ar_JO":
      return u"阿拉伯语"
   elif lang == "bg_BG":
      return u"保加利亚语-保加利亚"
   elif lang == "bs_BA":
      return u"意大利语"
   elif lang == "de_AT":
      return u"德语"
   elif lang == "de_CH":
      return u"德语"
   elif lang == "de":
      return u"马来语"
   elif lang == "el_CY":
      return u"印地语"
   elif lang == "en_IE":
      return u"英语"
   elif lang == "en_NZ":
      return u"西班牙语"
   elif lang == "en_TH":
      return u"英语"
   elif lang == "fr_CH":
      return u"法语"
   elif lang == "id_ID":
      return u"印尼爪哇语"
   elif lang == "la":
      return u"英语"
   elif lang == "lb_LU":
      return u"英文"
   elif lang == "mt_MT":
      return u"英文"
   elif lang == "nl_BE":
      return u"英文"
   elif lang == "sk_SK":
      return u"斯洛伐克语"
   elif lang == "ro_RO":
      return u"罗马尼亚语"
   elif lang == "th_TH":
      return u"泰语"
   elif lang == "tw_GH":
      return u"英文"
   elif lang == "he_IL":
      return u"希伯来语"
   elif lang == "en":
      return u"英语"
   elif lang == "es":
      return u"英语"
   elif lang == "cs_CZ":
      return u"捷克语-捷克"
   elif lang == "dk_DK":
      return u"丹麦语-丹麦"
   elif lang == "el_GR":
      return u"希腊语-希腊"
   elif lang == "en_AU":
      return u"英语-澳大利亚"
   elif lang == "en_CA":
      return u"英语-加拿大"
   elif lang == "es_ES":
      return u"西班牙语-西班牙"
   elif lang == "es_MX":
      return u"西班牙语-墨西哥"
   elif lang == "et_EE":
      return u"爱沙尼亚语-爱沙尼亚"
   elif lang == "fa_IR":
      return u"波斯语-伊朗"
   elif lang == "fi_FI":
      return u"芬兰语-芬兰"
   elif lang == "hi_IN":
      return u"印地语-印度"
   elif lang == "hr_HR":
      return u"克罗地亚语-克罗地亚"
   elif lang == "hu_HU":
      return u"匈牙利语-匈牙利"
   elif lang == "hy_AM":
      return u"亚美尼亚语-亚美尼亚"
   elif lang == "it_IT":
      return u"意大利语-意大利"
   elif lang == "ka_GE":
      return u"格鲁吉亚语-格鲁吉亚"
   elif lang == "lt_LT":
      return u"立陶宛语-立陶宛"
   elif lang == "lv_LV":
      return u"拉脱维亚语-拉脱维亚"
   elif lang == "ne_NP":
      return u"尼泊尔语-尼泊尔"
   elif lang == "nl_NL":
      return u"德语-荷兰"
   elif lang == "no_NO":
      return u"挪威语-挪威"
   elif lang == "pl_PL":
      return u"波兰语-波兰"
   elif lang == "pt_BR":
      return u"葡萄牙语-巴西"
   elif lang == "pt_PT":
      return u"葡萄牙语-葡萄牙"
   elif lang == "ru_RU":
      return u"俄语-俄罗斯"
   elif lang == "sl_SI":
      return u"斯诺文尼亚语-斯诺文尼亚"
   elif lang == "sv_SE":
      return u"瑞典语-瑞典"
   elif lang == "tr_TR":
      return u"土耳其语-土耳其"
   elif lang == "uk_UA":
      return u"乌克兰语-乌克兰"
   else:
      return u"未知语种"

def matchfaker(lang):
   if lang == "zh_CN":
      faker = Faker(lang)
      return faker.text()
   elif lang == "zh_TW":
      faker = Faker(lang)
      return faker.text()
   elif lang == "en_US":
      faker = Faker(lang)
      return faker.text()
   elif lang == "en_GB":
      faker = Faker(lang)
      return faker.text()
   elif lang == "de_DE":
      faker = Faker(lang)
      return faker.address()
   elif lang == "ja_JP":
      faker = Faker(lang)
      return faker.text()
   elif lang == "ko_KR":
      faker = Faker(lang)
      return faker.address()
   elif lang == "fr_FR":
      faker = Faker(lang)
      return faker.address()
   elif lang == "ar_AA":
      faker = Faker(lang)
      return faker.text()
   elif lang == "ar_EG":
      faker = Faker(lang)
      return faker.address()
   elif lang == "ar_PS":
      faker = Faker(lang)
      return faker.address()
   elif lang == "ar_SA":
      faker = Faker(lang)
      return faker.address()
   elif lang == "ar_JO":
      faker = Faker(lang)
      return faker.text()
   elif lang == "bg_BG":
      faker = Faker(lang)
      return faker.address()
   elif lang == "bs_BA":
      faker = Faker(lang)
      return faker.address()
   elif lang == "de_AT":
      faker = Faker(lang)
      return faker.address()
   elif lang == "de_CH":
      faker = Faker(lang)
      return faker.address()
   elif lang == "de":
      faker = Faker(lang)
      return faker.text()
   elif lang == "el_CY":
      faker = Faker(lang)
      return faker.text()
   elif lang == "en_IE":
      faker = Faker(lang)
      return faker.text()
   elif lang == "en_NZ":
      faker = Faker(lang)
      return faker.address()
   elif lang == "en_TH":
      faker = Faker(lang)
      return faker.text()
   elif lang == "fr_CH":
      faker = Faker(lang)
      return faker.address()
   elif lang == "id_ID":
      faker = Faker(lang)
      return faker.address()
   elif lang == "la":
      faker = Faker(lang)
      return faker.text()
   elif lang == "lb_LU":
      faker = Faker(lang)
      return faker.text()
   elif lang == "mt_MT":
      faker = Faker(lang)
      return faker.text()
   elif lang == "nl_BE":
      faker = Faker(lang)
      return faker.text()
   elif lang == "sk_SK":
      faker = Faker(lang)
      return faker.address()
   elif lang == "ro_RO":
      faker = Faker(lang)
      return faker.address()
   elif lang == "th_TH":
      faker = Faker(lang)
      return faker.text()
   elif lang == "tw_GH":
      faker = Faker(lang)
      return faker.text()
   elif lang == "he_IL":
      faker = Faker(lang)
      return faker.text()
   elif lang == "en":
      faker = Faker(lang)
      return faker.text()
   elif lang == "es":
      faker = Faker(lang)
      return faker.text()
   elif lang == "cs_CZ":
      faker = Faker(lang)
      return faker.address()
   elif lang == "dk_DK":
      faker = Faker(lang)
      return faker.address()
   elif lang == "el_GR":
      faker = Faker(lang)
      return faker.text()
   elif lang == "en_AU":
      faker = Faker(lang)
      return faker.text()
   elif lang == "en_CA":
      faker = Faker(lang)
      return faker.text()
   elif lang == "es_ES":
      faker = Faker(lang)
      return faker.address()
   elif lang == "es_MX":
      faker = Faker(lang)
      return faker.address()
   elif lang == "et_EE":
      faker = Faker(lang)
      return faker.address()
   elif lang == "fa_IR":
      faker = Faker(lang)
      return faker.address()
   elif lang == "fi_FI":
      faker = Faker(lang)
      return faker.address()
   elif lang == "hi_IN":
      faker = Faker(lang)
      return faker.address()
   elif lang == "hr_HR":
      faker = Faker(lang)
      return faker.address()
   elif lang == "hu_HU":
      faker = Faker(lang)
      return faker.address()
   elif lang == "hy_AM":
      faker = Faker(lang)
      return faker.text()
   elif lang == "it_IT":
      faker = Faker(lang)
      return faker.address()
   elif lang == "ka_GE":
      faker = Faker(lang)
      return faker.address()
   elif lang == "lt_LT":
      faker = Faker(lang)
      return faker.address()
   elif lang == "lv_LV":
      faker = Faker(lang)
      return faker.address()
   elif lang == "ne_NP":
      faker = Faker(lang)
      return faker.address()
   elif lang == "nl_NL":
      faker = Faker(lang)
      return faker.address()
   elif lang == "no_NO":
      faker = Faker(lang)
      return faker.address()
   elif lang == "pl_PL":
      faker = Faker(lang)
      return faker.text()
   elif lang == "pt_BR":
      faker = Faker(lang)
      return faker.address()
   elif lang == "pt_PT":
      faker = Faker(lang)
      return faker.address()
   elif lang == "ru_RU":
      faker = Faker(lang)
      return faker.text()
   elif lang == "sl_SI":
      faker = Faker(lang)
      return faker.address()
   elif lang == "sv_SE":
      faker = Faker(lang)
      return faker.address()
   elif lang == "tr_TR":
      faker = Faker(lang)
      return faker.address()
   elif lang == "uk_UA":
      faker = Faker(lang)
      return faker.address()
   else:
      return u"未知语种"

def CreateWord(lang):
   faker = Faker(lang)
   filename = "test-%s-%s-%s.docx" % (faker.name(),lang,matchlanguage(lang))
   #filename = faker.file_name(category=None, extension="docx")
   print filename
   path = "Result\\multilang"
   makedir(path)
   i = 0
   z = 0
   document = Document()
   s = random.randint(0,10)
   while z < s:
      text1 = matchfaker(lang)
      time.sleep(1)
      document.add_paragraph(text1)
      z = z + 1
   text2 = matchfaker(lang)
   time.sleep(0.5)
   document.add_paragraph(u'生涯规划')
   document.add_paragraph(text2)
   document.save(u"%s\\%s" % (path,filename))
   return True

def run():
   for lang in lang_list:
      CreateWord(lang)

if __name__ == "__main__":
   run()
